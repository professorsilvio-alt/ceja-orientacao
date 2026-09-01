import os
import io
from django.conf import settings

# Importações com fallback seguro para não quebrar inicialização do Django no servidor
try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False

try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    try:
        import PyPDF2 as pypdf
        PYPDF_AVAILABLE = True
    except ImportError:
        PYPDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False


def dividir_em_fragmentos(texto: str, tamanho_max_caracteres: int = 1200, sobreposicao: int = 150) -> list[str]:
    """
    Divide um texto longo em fragmentos (chunks) com sobreposição para contexto no RAG.
    """
    if not texto:
        return []
    
    paragrafos = texto.split('\n')
    fragmentos = []
    chunk_atual = ""

    for p in paragrafos:
        p = p.strip()
        if not p:
            continue
        if len(chunk_atual) + len(p) + 1 <= tamanho_max_caracteres:
            chunk_atual += ("\n" if chunk_atual else "") + p
        else:
            if chunk_atual:
                fragmentos.append(chunk_atual.strip())
            # Iniciar novo chunk mantendo um pouco do contexto se possível
            if len(p) > tamanho_max_caracteres:
                for i in range(0, len(p), tamanho_max_caracteres - sobreposicao):
                    fragmentos.append(p[i:i + tamanho_max_caracteres].strip())
                chunk_atual = ""
            else:
                chunk_atual = p

    if chunk_atual:
        fragmentos.append(chunk_atual.strip())

    return [f for f in fragmentos if f]


def extrair_texto_pdf(caminho_arquivo: str) -> str:
    """Extrai texto de um PDF usando PyMuPDF, pypdf ou PyPDF2 com fallback."""
    textos = []
    
    if FITZ_AVAILABLE:
        try:
            doc = fitz.open(caminho_arquivo)
            for num_pagina in range(len(doc)):
                pagina = doc[num_pagina]
                txt = pagina.get_text()
                if txt.strip():
                    textos.append(f"--- Página {num_pagina + 1} ---\n{txt.strip()}")
            doc.close()
            return "\n\n".join(textos)
        except Exception as e:
            pass

    if PYPDF_AVAILABLE:
        try:
            reader = pypdf.PdfReader(caminho_arquivo)
            for num_pagina, page in enumerate(reader.pages):
                txt = page.extract_text()
                if txt and txt.strip():
                    textos.append(f"--- Página {num_pagina + 1} ---\n{txt.strip()}")
            return "\n\n".join(textos)
        except Exception as e:
            return f"[Erro na extração de PDF via pypdf: {e}]"

    return "[Instale 'pymupdf' ou 'pypdf' para extração automática de arquivos PDF no servidor]"


def extrair_texto_docx(caminho_arquivo: str) -> str:
    """Extrai texto de um arquivo Word (.docx)."""
    if not DOCX_AVAILABLE:
        return "[Instale 'python-docx' para extração automática de arquivos Word]"

    try:
        doc = docx.Document(caminho_arquivo)
        paragrafos = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Extrair tabelas se houver
        tabelas_texto = []
        for tabela in doc.tables:
            linhas = []
            for row in tabela.rows:
                celulas = [c.text.strip() for c in row.cells]
                linhas.append(" | ".join(celulas))
            if linhas:
                tabelas_texto.append("Tabela:\n" + "\n".join(linhas))

        resultado = "\n\n".join(paragrafos)
        if tabelas_texto:
            resultado += "\n\n" + "\n\n".join(tabelas_texto)
        return resultado
    except Exception as e:
        return f"[Erro na extração de Word: {e}]"


def extrair_texto_planilha(caminho_arquivo: str) -> str:
    """Extrai dados estruturados de planilhas Excel (.xlsx, .xls) ou CSV."""
    if not PANDAS_AVAILABLE:
        # Fallback para leitura de texto puro se for CSV
        if caminho_arquivo.lower().endswith('.csv'):
            return extrair_texto_txt(caminho_arquivo)
        return "[Instale 'pandas' e 'openpyxl' para leitura automática de planilhas]"

    try:
        ext = os.path.splitext(caminho_arquivo)[1].lower()
        if ext == '.csv':
            df = pd.read_csv(caminho_arquivo)
            return f"Planilha CSV:\n" + df.to_string(index=False)
        else:
            excel = pd.ExcelFile(caminho_arquivo)
            abas_texto = []
            for sheet_name in excel.sheet_names:
                df = pd.read_excel(excel, sheet_name=sheet_name)
                abas_texto.append(f"=== Aba: {sheet_name} ===\n" + df.to_string(index=False))
            return "\n\n".join(abas_texto)
    except Exception as e:
        return f"[Erro na extração de Planilha: {e}]"


def extrair_texto_txt(caminho_arquivo: str) -> str:
    """Lê arquivo de texto puro com fallback de codificações UTF-8 e Latin-1."""
    encodings = ['utf-8', 'latin-1', 'cp1252']
    for enc in encodings:
        try:
            with open(caminho_arquivo, 'r', encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    return "[Erro na decodificação do arquivo de texto]"


def processar_arquivo_documento(doc_instance) -> str:
    """
    Processa o arquivo vinculado a uma instância de DocumentoCerebro
    e retorna o texto extraído correspondente.
    """
    if not doc_instance.arquivo:
        return doc_instance.conteudo_extraido or ""

    caminho = doc_instance.arquivo.path
    if not os.path.exists(caminho):
        return doc_instance.conteudo_extraido or ""

    ext = os.path.splitext(caminho)[1].lower()
    
    if ext == '.pdf':
        doc_instance.tipo_arquivo = 'pdf'
        return extrair_texto_pdf(caminho)
    elif ext in ['.docx', '.doc']:
        doc_instance.tipo_arquivo = 'docx'
        return extrair_texto_docx(caminho)
    elif ext in ['.xlsx', '.xls', '.csv']:
        doc_instance.tipo_arquivo = 'xlsx'
        return extrair_texto_planilha(caminho)
    elif ext in ['.txt', '.log']:
        doc_instance.tipo_arquivo = 'texto'
        return extrair_texto_txt(caminho)
    elif ext in ['.mp3', '.wav', '.m4a', '.ogg', '.webm']:
        doc_instance.tipo_arquivo = 'audio'
        return f"[Áudio gravado/enviado: {os.path.basename(caminho)} - Processamento multimídia ativado]"
    elif ext in ['.png', '.jpg', '.jpeg', '.webp']:
        doc_instance.tipo_arquivo = 'imagem'
        return f"[Imagem/Print anexado: {os.path.basename(caminho)} - Análise visual ativada]"
    else:
        return extrair_texto_txt(caminho)
